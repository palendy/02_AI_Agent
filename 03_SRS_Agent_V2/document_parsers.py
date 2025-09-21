"""
Document parsers for various file formats (PDF, DOCX, TXT, MD).
Provides unified interface for extracting text content from different document types.
"""

import os
import re
import logging
from typing import Dict, List, Optional, Union, Any
from pathlib import Path
from abc import ABC, abstractmethod

# Document processing libraries
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import markdown
except ImportError:
    markdown = None

try:
    import chardet
except ImportError:
    chardet = None

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

from config import get_config

logger = logging.getLogger(__name__)


class DocumentParser(ABC):
    """Abstract base class for document parsers."""
    
    def __init__(self, config=None):
        self.config = config or get_config()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.processing.chunk_size,
            chunk_overlap=self.config.processing.chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    @abstractmethod
    def parse(self, file_path: str) -> List[Document]:
        """Parse document and return list of Document objects."""
        pass
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """Extract raw text content from document."""
        pass
    
    def validate_file(self, file_path: str) -> bool:
        """Validate if file exists and is readable."""
        try:
            path = Path(file_path)
            if not path.exists():
                logger.error(f"File does not exist: {file_path}")
                return False
            
            if not path.is_file():
                logger.error(f"Path is not a file: {file_path}")
                return False
            
            # Check file size
            size_mb = path.stat().st_size / (1024 * 1024)
            if size_mb > self.config.processing.max_file_size_mb:
                logger.error(f"File too large: {size_mb}MB > {self.config.processing.max_file_size_mb}MB")
                return False
            
            return True
        except Exception as e:
            logger.error(f"Error validating file {file_path}: {str(e)}")
            return False
    
    def create_documents(self, text: str, file_path: str, metadata: Dict[str, Any] = None) -> List[Document]:
        """Create Document objects from text with metadata."""
        if not text.strip():
            logger.warning(f"No text content extracted from {file_path}")
            return []
        
        base_metadata = {
            "source": file_path,
            "file_name": Path(file_path).name,
            "file_type": Path(file_path).suffix.lower().lstrip('.'),
            "file_size": os.path.getsize(file_path),
            "total_length": len(text)
        }
        
        if metadata:
            base_metadata.update(metadata)
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        
        documents = []
        for i, chunk in enumerate(chunks):
            doc_metadata = base_metadata.copy()
            doc_metadata.update({
                "chunk_index": i,
                "chunk_length": len(chunk)
            })
            documents.append(Document(page_content=chunk, metadata=doc_metadata))
        
        logger.info(f"Created {len(documents)} documents from {file_path}")
        return documents


class PDFParser(DocumentParser):
    """Parser for PDF documents."""
    
    def __init__(self, config=None):
        super().__init__(config)
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install PyPDF2")
    
    def extract_text(self, file_path: str) -> str:
        """Extract text content from PDF file."""
        if not self.validate_file(file_path):
            return ""
        
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
            
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            return ""
    
    def parse(self, file_path: str) -> List[Document]:
        """Parse PDF and return Document objects."""
        text = self.extract_text(file_path)
        if not text:
            return []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata = {
                    "pages": len(pdf_reader.pages),
                    "title": getattr(pdf_reader.metadata, 'title', None) if pdf_reader.metadata else None,
                    "author": getattr(pdf_reader.metadata, 'author', None) if pdf_reader.metadata else None,
                    "creator": getattr(pdf_reader.metadata, 'creator', None) if pdf_reader.metadata else None,
                }
        except Exception as e:
            logger.warning(f"Could not extract PDF metadata: {str(e)}")
            metadata = {}
        
        return self.create_documents(text, file_path, metadata)


class DOCXParser(DocumentParser):
    """Parser for DOCX documents."""
    
    def __init__(self, config=None):
        super().__init__(config)
        if DocxDocument is None:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
    
    def extract_text(self, file_path: str) -> str:
        """Extract text content from DOCX file."""
        if not self.validate_file(file_path):
            return ""
        
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
        
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            return ""
    
    def parse(self, file_path: str) -> List[Document]:
        """Parse DOCX and return Document objects."""
        text = self.extract_text(file_path)
        if not text:
            return []
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title,
                "author": core_props.author,
                "subject": core_props.subject,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
                "pages": len(doc.paragraphs),  # Approximate page count
                "tables": len(doc.tables)
            }
        except Exception as e:
            logger.warning(f"Could not extract DOCX metadata: {str(e)}")
            metadata = {}
        
        return self.create_documents(text, file_path, metadata)


class TXTParser(DocumentParser):
    """Parser for plain text documents."""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text content from TXT file."""
        if not self.validate_file(file_path):
            return ""
        
        try:
            # Detect encoding
            encoding = 'utf-8'
            if chardet:
                with open(file_path, 'rb') as file:
                    raw_data = file.read()
                    result = chardet.detect(raw_data)
                    if result['encoding']:
                        encoding = result['encoding']
            
            # Read file with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='replace') as file:
                return file.read()
        
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {str(e)}")
            try:
                # Fallback to different encodings
                for enc in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        with open(file_path, 'r', encoding=enc, errors='replace') as file:
                            logger.info(f"Successfully read {file_path} with {enc} encoding")
                            return file.read()
                    except:
                        continue
            except:
                pass
            return ""
    
    def parse(self, file_path: str) -> List[Document]:
        """Parse TXT and return Document objects."""
        text = self.extract_text(file_path)
        if not text:
            return []
        
        # Basic text file metadata
        metadata = {
            "lines": len(text.split('\n')),
            "words": len(text.split()),
            "characters": len(text)
        }
        
        return self.create_documents(text, file_path, metadata)


class MDParser(DocumentParser):
    """Parser for Markdown documents."""
    
    def __init__(self, config=None):
        super().__init__(config)
        if markdown is None:
            raise ImportError("markdown is required for MD parsing. Install with: pip install markdown")
    
    def extract_text(self, file_path: str) -> str:
        """Extract text content from Markdown file."""
        if not self.validate_file(file_path):
            return ""
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                content = file.read()
            
            # Convert markdown to HTML then strip HTML tags for plain text
            html_content = markdown.markdown(content)
            
            # Remove HTML tags to get plain text
            clean_text = re.sub('<[^<]+?>', '', html_content)
            
            # Also preserve the original markdown for structure
            return f"=== ORIGINAL MARKDOWN ===\n{content}\n\n=== PLAIN TEXT ===\n{clean_text}"
        
        except Exception as e:
            logger.error(f"Error parsing MD {file_path}: {str(e)}")
            return ""
    
    def parse(self, file_path: str) -> List[Document]:
        """Parse Markdown and return Document objects."""
        text = self.extract_text(file_path)
        if not text:
            return []
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                content = file.read()
            
            # Extract markdown-specific metadata
            headers = re.findall(r'^#+\s+(.+)$', content, re.MULTILINE)
            code_blocks = re.findall(r'```[\s\S]*?```', content)
            links = re.findall(r'\[([^\]]+)\]\(([^)]+)\)', content)
            
            metadata = {
                "headers": headers,
                "header_count": len(headers),
                "code_blocks": len(code_blocks),
                "links": len(links),
                "lines": len(content.split('\n'))
            }
        except Exception as e:
            logger.warning(f"Could not extract MD metadata: {str(e)}")
            metadata = {}
        
        return self.create_documents(text, file_path, metadata)


class DocumentParserFactory:
    """Factory class for creating appropriate document parsers."""
    
    _parsers = {
        '.pdf': PDFParser,
        '.docx': DOCXParser,
        '.doc': DOCXParser,  # Treat .doc same as .docx (may need additional library for true .doc support)
        '.txt': TXTParser,
        '.text': TXTParser,
        '.md': MDParser,
        '.markdown': MDParser
    }
    
    @classmethod
    def create_parser(cls, file_path: str, config=None) -> Optional[DocumentParser]:
        """Create appropriate parser for given file type."""
        suffix = Path(file_path).suffix.lower()
        
        if suffix not in cls._parsers:
            logger.error(f"Unsupported file format: {suffix}")
            return None
        
        try:
            parser_class = cls._parsers[suffix]
            return parser_class(config)
        except Exception as e:
            logger.error(f"Error creating parser for {suffix}: {str(e)}")
            return None
    
    @classmethod
    def get_supported_extensions(cls) -> List[str]:
        """Get list of supported file extensions."""
        return list(cls._parsers.keys())
    
    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """Check if file type is supported."""
        suffix = Path(file_path).suffix.lower()
        return suffix in cls._parsers


class MultiDocumentParser:
    """Parser that can handle multiple documents of different types."""
    
    def __init__(self, config=None):
        self.config = config or get_config()
        self.factory = DocumentParserFactory()
    
    def parse_documents(self, file_paths: List[str]) -> Dict[str, List[Document]]:
        """Parse multiple documents and return organized results."""
        results = {}
        
        for file_path in file_paths:
            try:
                parser = self.factory.create_parser(file_path, self.config)
                if parser:
                    documents = parser.parse(file_path)
                    results[file_path] = documents
                    logger.info(f"Successfully parsed {file_path}: {len(documents)} documents")
                else:
                    logger.warning(f"No parser available for {file_path}")
                    results[file_path] = []
            except Exception as e:
                logger.error(f"Error parsing {file_path}: {str(e)}")
                results[file_path] = []
        
        return results
    
    def parse_single_document(self, file_path: str) -> List[Document]:
        """Parse a single document and return Document objects."""
        parser = self.factory.create_parser(file_path, self.config)
        if parser:
            return parser.parse(file_path)
        return []
    
    def extract_all_text(self, file_paths: List[str]) -> Dict[str, str]:
        """Extract text from multiple documents."""
        results = {}
        
        for file_path in file_paths:
            try:
                parser = self.factory.create_parser(file_path, self.config)
                if parser:
                    text = parser.extract_text(file_path)
                    results[file_path] = text
                    logger.info(f"Successfully extracted text from {file_path}: {len(text)} characters")
                else:
                    logger.warning(f"No parser available for {file_path}")
                    results[file_path] = ""
            except Exception as e:
                logger.error(f"Error extracting text from {file_path}: {str(e)}")
                results[file_path] = ""
        
        return results


# Utility functions
def detect_document_type(file_path: str) -> Optional[str]:
    """Detect document type from file extension."""
    suffix = Path(file_path).suffix.lower()
    type_mapping = {
        '.pdf': 'PDF',
        '.docx': 'DOCX',
        '.doc': 'DOC',
        '.txt': 'Text',
        '.text': 'Text',
        '.md': 'Markdown',
        '.markdown': 'Markdown'
    }
    return type_mapping.get(suffix)


def get_document_info(file_path: str) -> Dict[str, Any]:
    """Get basic information about a document file."""
    try:
        path = Path(file_path)
        if not path.exists():
            return {"error": "File does not exist"}
        
        stat = path.stat()
        return {
            "name": path.name,
            "size": stat.st_size,
            "size_mb": round(stat.st_size / (1024 * 1024), 2),
            "type": detect_document_type(file_path),
            "extension": path.suffix.lower(),
            "supported": DocumentParserFactory.is_supported(file_path),
            "modified": stat.st_mtime
        }
    except Exception as e:
        return {"error": str(e)}


def batch_parse_directory(directory_path: str, config=None) -> Dict[str, List[Document]]:
    """Parse all supported documents in a directory."""
    directory = Path(directory_path)
    if not directory.exists() or not directory.is_dir():
        logger.error(f"Directory does not exist or is not a directory: {directory_path}")
        return {}
    
    # Find all supported files
    supported_files = []
    for ext in DocumentParserFactory.get_supported_extensions():
        pattern = f"*{ext}"
        supported_files.extend(directory.glob(pattern))
    
    # Convert to strings and parse
    file_paths = [str(f) for f in supported_files]
    parser = MultiDocumentParser(config)
    
    logger.info(f"Found {len(file_paths)} supported files in {directory_path}")
    return parser.parse_documents(file_paths)